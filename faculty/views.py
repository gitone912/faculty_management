# In faculty_management/views.py
from django.shortcuts import render, redirect
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.urls import reverse
from django.db.models import Sum
from django.utils.timezone import now
from .models import User, Subject, Course, TeachingRecord

from django.shortcuts import render, redirect
from django.contrib.auth.decorators import login_required
from django.urls import reverse
from django.contrib import messages
from django.utils import timezone
from .models import TeachingRecord, Subject, Course
from datetime import datetime
from django.core.paginator import Paginator

from django.shortcuts import render, redirect
from django.contrib.auth.decorators import login_required
from django.urls import reverse
from django.contrib import messages
from .models import TeachingRecord, Subject, Course
from datetime import datetime

@login_required
def record_teaching_activity(request):
    """
    View for faculty to record their teaching activities.
    """
    # Handle POST request for recording teaching activity
    if request.method == 'POST':
        # Retrieve form data
        subject_id = request.POST.get('subject')
        course_id = request.POST.get('course')
        topic_taught = request.POST.get('topic_taught')
        description = request.POST.get('description')
        
        # Retrieve start and end time from the form
        start_time_str = request.POST.get('start_time')
        end_time_str = request.POST.get('end_time')
        
        # Convert the start and end times to datetime objects using the correct method
        start_time = datetime.datetime.fromisoformat(start_time_str)
        end_time = datetime.datetime.fromisoformat(end_time_str)
        
        # Retrieve subject and course instances
        subject = Subject.objects.get(id=subject_id)
        course = Course.objects.get(id=course_id)
        
        # Check for overlapping time frames with existing teaching records
        overlapping_records = TeachingRecord.objects.filter(
            faculty=request.user,
            start_time__lt=end_time,  # Records where start_time is before end_time
            end_time__gt=start_time  # Records where end_time is after start_time
        )

        # If there is an overlap, show an error message and do not save the record
        if overlapping_records.exists():
            messages.error(request, 'The specified time frame overlaps with another class you are teaching. Please choose a different time frame.')
            return redirect(reverse('record_teaching_activity'))

        # If no overlap, create a new TeachingRecord instance
        teaching_record = TeachingRecord(
            faculty=request.user,
            subject=subject,
            course=course,
            topic_taught=topic_taught,
            description=description,
            start_time=start_time,
            end_time=end_time,
        )
        
        # Save the TeachingRecord instance
        teaching_record.save()
        messages.success(request, 'Teaching activity recorded successfully.')
        return redirect(reverse('faculty_dashboard'))

    # Retrieve available subjects and courses for the faculty
    subjects = Subject.objects.all()
    courses = Course.objects.all()

    # Render the template with the context
    return render(request, 'record_teaching_activity.html', {'subjects': subjects, 'courses': courses})


@login_required
def view_salary_information(request):
    """
    View for faculty to view their daily and monthly salary.
    """
    # if not request.user.is_faculty:
    #     return redirect('home')  # Ensure the user is a faculty member

    faculty_id = request.user.id

    # Calculate daily and monthly salary for the faculty
    today = now().date()
    start_of_month = today.replace(day=1)

    daily_salary = TeachingRecord.objects.filter(
        faculty_id=faculty_id,
        date_taught=today
    ).aggregate(total_salary=Sum('salary'))['total_salary'] or 0.00

    monthly_salary = TeachingRecord.objects.filter(
        faculty_id=faculty_id,
        date_taught__gte=start_of_month,
        date_taught__lte=today
    ).aggregate(total_salary=Sum('salary'))['total_salary'] or 0.00

    context = {
        'daily_salary': daily_salary,
        'monthly_salary': monthly_salary,
    }

    return render(request, 'salary_report.html', context)



# In faculty_management/views.py
from django.shortcuts import render, redirect
from django.contrib.auth import authenticate, login
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.urls import reverse
from django.db.models import Sum
from django.utils.timezone import now
from .models import User, Subject, Course, TeachingRecord

from django.contrib.auth import authenticate, login
from django.shortcuts import render, redirect
from django.urls import reverse
from django.contrib import messages

def faculty_login(request):
    """
    View for faculty login.
    """
    if request.method == 'POST':
        # Retrieve login credentials from the form
        username = request.POST.get('username')
        password = request.POST.get('password')

        # Authenticate the user
        user = authenticate(request, username=username, password=password)

        if user is not None:
            # Log the user in
            login(request, user)
            
            # Check if the user is a superuser
            if user.is_superuser:
                # Redirect to the superuser dashboard
                return redirect(reverse('superuser_dashboard'))
            else:
                # Redirect to a faculty dashboard or home page
                return redirect(reverse('faculty_dashboard'))
        else:
            # Authentication failed
            messages.error(request, 'Invalid login credentials.')
            return redirect(reverse('faculty_login'))

    # Render the login page
    return render(request, 'faculty_login.html')




from django.shortcuts import render
from django.contrib.auth.decorators import login_required
from .models import TeachingRecord

@login_required
def faculty_dashboard(request):
    
    user = request.user
    if request.user.is_superuser:
        # Redirect superusers to the superuser dashboard
        return redirect(reverse('superuser_dashboard'))
    # Filter records for the current user and order by start_time in descending order
    teaching_records = TeachingRecord.objects.filter(faculty=user).order_by('-start_time')
    
    # Calculate total classes and salary
    total_classes = teaching_records.count()
    total_salary = sum(record.salary for record in teaching_records)
    
    # Calculate total days taught based on unique dates from start_time
    total_days = len(set(record.start_time.date() for record in teaching_records))
    
    # Calculate total hours taught by summing the durations of each record
    total_hours = sum((record.end_time - record.start_time).total_seconds() / 3600 for record in teaching_records)
    
    # Implement pagination
    paginator = Paginator(teaching_records, 10)  # Set the number of records per page
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    
    # Pass data to template
    context = {
        'teaching_records': page_obj,
        'total_classes': total_classes,
        'total_salary': total_salary,
        'total_days': total_days,
        'total_hours': round(total_hours),
        'paginator': paginator,  # Optional: pass the paginator object if needed
    }
    return render(request, 'faculty_dashboard.html', context)

from django.contrib.auth.decorators import login_required, user_passes_test

@login_required
@user_passes_test(lambda user: user.is_superuser)
def superuser_dashboard(request):
    # Fetch all teaching records sorted by latest entry (start_time in descending order)
    teaching_records = TeachingRecord.objects.all().order_by('-start_time')
    
    # Calculate total classes, days, hours, and salary
    total_classes = teaching_records.count()
    total_salary = sum(record.salary for record in teaching_records)
    total_days = len(set(record.start_time.date() for record in teaching_records if record.start_time is not None))
    total_hours = sum((record.end_time - record.start_time).total_seconds() / 3600 for record in teaching_records if record.start_time is not None and record.end_time is not None)
    
    # Implement pagination
    paginator = Paginator(teaching_records, 10)  # Set the number of records per page
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    
    # Pass data to the template
    context = {
        'teaching_records': page_obj,
        'total_classes': total_classes,
        'total_salary': total_salary,
        'total_days': total_days,
        'total_hours': round(total_hours),
        'paginator': paginator,  # Optional: pass the paginator object if needed
    }
    
    return render(request, 'superuser_dashboard.html', context)

from django.http import HttpResponse
import csv
from docx import Document
from django.utils import timezone
from .models import TeachingRecord

@login_required
def download_csv(request):
    # Create an HTTP response with the content type for CSV
    response = HttpResponse(content_type='text/csv')
    # Specify the CSV file name for the HTTP response header
    response['Content-Disposition'] = f'attachment; filename="teaching_records_{timezone.now().date()}.csv"'

    # Create a CSV writer
    writer = csv.writer(response)
    # Write the header row
    writer.writerow(['Subject', 'Course', 'Topic Taught', 'Date Taught', 'Time Slot', 'Salary'])

    # Retrieve start_time and end_time from query parameters
    start_time_str = request.GET.get('start_time')
    end_time_str = request.GET.get('end_time')

    # Initialize variables for start_time and end_time
    start_time = None
    end_time = None
    
    if start_time_str and end_time_str:
        # Convert query parameters to datetime objects
        start_time = datetime.datetime.strptime(start_time_str, '%Y-%m-%d')
        end_time = datetime.datetime.strptime(end_time_str, '%Y-%m-%d')
    
    # Filter teaching records based on the date range
    if start_time and end_time:
        records = TeachingRecord.objects.filter(faculty=request.user, start_time__date__range=[start_time, end_time])
    else:
        # If no date range is specified, default to all records for the user
        records = TeachingRecord.objects.filter(faculty=request.user)

    # Write the data rows
    for record in records:
        writer.writerow([
            record.subject.name,
            record.course.name,
            record.topic_taught,
            record.start_time.strftime('%Y-%m-%d'),
            f"{record.start_time.strftime('%I:%M %p')} - {record.end_time.strftime('%I:%M %p')}",
            record.salary
        ])

    return response

@login_required
def download_doc(request):
    # Get the start_time and end_time from the query parameters
    start_time_str = request.GET.get('start_time')
    end_time_str = request.GET.get('end_time')
    print(start_time_str)
    print(end_time_str)
    # Initialize variables for start_time and end_time
    start_time = None
    end_time = None
    
    if start_time_str and end_time_str:
        # Parse start time and end time from query parameters
        start_time = datetime.datetime.strptime(start_time_str, '%Y-%m-%d')
        end_time = datetime.datetime.strptime(end_time_str, '%Y-%m-%d')
    
    # Filter teaching records based on the date range
    if start_time and end_time:
        records = TeachingRecord.objects.filter(faculty=request.user, start_time__date__range=[start_time, end_time])
    else:
        # If no date range is specified, default to all records for the user
        records = TeachingRecord.objects.filter(faculty=request.user)
    
    # Create a new Word Document
    doc = Document()
    # Add a title to the document
    doc.add_heading('Teaching Records', level=1)
    
    # Add a table to the document
    table = doc.add_table(rows=1, cols=6)
    table.autofit = True
    
    # Add a header row
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Subject'
    header_cells[1].text = 'Course'
    header_cells[2].text = 'Topic Taught'
    header_cells[3].text = 'Date Taught'
    header_cells[4].text = 'Time Slot'
    header_cells[5].text = 'Salary'
    
    # Add rows to the table for each record
    for record in records:
        row_cells = table.add_row().cells
        row_cells[0].text = record.subject.name
        row_cells[1].text = record.course.name
        row_cells[2].text = record.topic_taught
        row_cells[3].text = record.start_time.strftime('%Y-%m-%d')
        row_cells[4].text = f"{record.start_time.strftime('%I:%M %p')} - {record.end_time.strftime('%I:%M %p')}"
        row_cells[5].text = f"₹{record.salary}"
    
    # Create an HTTP response for the DOCX file
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="teaching_records_{timezone.now().date()}.docx"'
    
    # Save the document to the response
    doc.save(response)

    return response


from django.contrib.auth import logout
from django.shortcuts import redirect

def faculty_logout(request):
    logout(request)
    return redirect('faculty_login')  # Redirect to the login page after logging out


from django.http import HttpResponse
from django.utils import timezone
import csv
from .models import TeachingRecord
import datetime

@login_required
@user_passes_test(lambda user: user.is_superuser)
def admin_csv(request):
    # Create an HTTP response with the content type for CSV
    response = HttpResponse(content_type='text/csv')
    # Specify the CSV file name for the HTTP response header
    response['Content-Disposition'] = f'attachment; filename="teaching_records_{timezone.now().date()}.csv"'

    # Create a CSV writer
    writer = csv.writer(response)
    # Write the header row
    writer.writerow(['Faculty', 'Subject', 'Course', 'Topic Taught', 'Date Taught', 'Time Slot', 'Salary'])

    # Retrieve start_time and end_time from query parameters
    start_time_str = request.GET.get('start_time')
    end_time_str = request.GET.get('end_time')
    
    # Initialize variables for start_time and end_time
    start_time = None
    end_time = None
    
    if start_time_str and end_time_str:
        # Convert query parameters to datetime objects
        start_time = datetime.datetime.strptime(start_time_str, '%Y-%m-%d')
        end_time = datetime.datetime.strptime(end_time_str, '%Y-%m-%d')
    
    # Filter teaching records based on the date range
    if start_time and end_time:
        records = TeachingRecord.objects.filter(start_time__date__range=[start_time, end_time])
    else:
        # If no date range is specified, default to all records
        records = TeachingRecord.objects.all()

    # Write the data rows
    total_salary = 0  # Initialize total salary

    for record in records:
        writer.writerow([
            record.faculty.get_full_name(),
            record.subject.name,
            record.course.name,
            record.topic_taught,
            record.start_time.strftime('%Y-%m-%d'),
            f"{record.start_time.strftime('%I:%M %p')} - {record.end_time.strftime('%I:%M %p')}",
            record.salary
        ])
        # Add record's salary to the total
        total_salary += record.salary

    # Add total salary row at the end
    writer.writerow(['', '', '', '', '', 'Total Salary:', total_salary])

    return response


from django.http import HttpResponse
from django.utils import timezone
from docx import Document
from .models import TeachingRecord
import datetime

@login_required
@user_passes_test(lambda user: user.is_superuser)
def admin_doc(request):
    # Retrieve start_time and end_time from query parameters
    start_time_str = request.GET.get('start_time')
    end_time_str = request.GET.get('end_time')

    # Initialize variables for start_time and end_time
    start_time = None
    end_time = None
    
    if start_time_str and end_time_str:
        # Convert query parameters to datetime objects
        start_time = datetime.datetime.strptime(start_time_str, '%Y-%m-%d')
        end_time = datetime.datetime.strptime(end_time_str, '%Y-%m-%d')
    
    # Filter teaching records based on the date range
    if start_time and end_time:
        records = TeachingRecord.objects.filter(start_time__date__range=[start_time, end_time])
    else:
        # If no date range is specified, default to all records
        records = TeachingRecord.objects.all()
    
    # Create a new Word Document
    doc = Document()
    # Add a title to the document
    doc.add_heading('Teaching Records', level=1)
    
    # Add a table to the document
    table = doc.add_table(rows=1, cols=7)
    table.autofit = True
    
    # Add a header row
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Faculty'
    header_cells[1].text = 'Subject'
    header_cells[2].text = 'Course'
    header_cells[3].text = 'Topic Taught'
    header_cells[4].text = 'Date Taught'
    header_cells[5].text = 'Time Slot'
    header_cells[6].text = 'Salary'
    
    # Add rows to the table for each record
    for record in records:
        row_cells = table.add_row().cells
        row_cells[0].text = record.faculty.get_full_name()
        row_cells[1].text = record.subject.name
        row_cells[2].text = record.course.name
        row_cells[3].text = record.topic_taught
        row_cells[4].text = record.start_time.strftime('%Y-%m-%d')
        row_cells[5].text = f"{record.start_time.strftime('%I:%M %p')} - {record.end_time.strftime('%I:%M %p')}"
        row_cells[6].text = f"₹{record.salary}"
    
    # Create an HTTP response for the DOCX file
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="teaching_records_{timezone.now().date()}.docx"'
    
    # Save the document to the response
    doc.save(response)

    return response


from django.shortcuts import render
from django.contrib.auth.decorators import login_required

@login_required
def profile_view(request):
    # Retrieve the current user
    user = request.user
    
    # Determine the user's role (Admin or Visiting Faculty)
    if user.is_superuser:
        role = "Admin"
        experience = None  # Don't show experience for superusers (admins)
    else:
        role = "Visiting Faculty"
        # Assume experience is stored in the user profile
        experience = user.profile.experience if hasattr(user, 'profile') and user.profile.experience is not None else 'N/A'

    # Context to pass to the template
    context = {
        'name': user.get_full_name(),
        'username': user.username,
        'email': user.email,
        'role': role,
        'experience': experience
    }

    # Render the profile template
    return render(request, 'profile.html', context)

def home(request):
    # Check if the user is authenticated (logged in)
    if request.user.is_authenticated:
        # If the user is logged in, redirect them to faculty_dashboard
        return redirect(reverse('faculty_dashboard'))
    
    # If the user is not logged in, render the homepage as usual
    return render(request, 'homepage.html')


from django.http import HttpResponse
from openpyxl import Workbook
from datetime import datetime
from .models import TeachingRecord, User

def download_excel(request):
    # Create a new workbook
    workbook = Workbook()
    
    # Fetch all faculty members (users with profiles)
    faculty_members = User.objects.filter(profile__isnull=False)
    
    # Retrieve start_date and end_date from query parameters
    start_date_str = request.GET.get('start_date')
    end_date_str = request.GET.get('end_date')
    
    # Initialize variables for start_date and end_date
    start_date = None
    end_date = None
    
    if start_date_str and end_date_str:
        # Convert query parameters to datetime.date objects
        try:
            start_date = datetime.strptime(start_date_str, '%Y-%m-%d').date()
            end_date = datetime.strptime(end_date_str, '%Y-%m-%d').date()
        except ValueError:
            return HttpResponse("Invalid date format. Please use YYYY-MM-DD.", status=400)
    
    # Loop through each faculty member
    for faculty in faculty_members:
        # Create a new worksheet for each faculty member
        sheet = workbook.create_sheet(title=faculty.username)
        
        # Add headers to the sheet
        headers = ['Subject', 'Course', 'Topic Taught', 'Start Time', 'End Time', 'Description', 'Salary']
        sheet.append(headers)
        
        # Filter teaching records based on the date range and the faculty
        teaching_records = TeachingRecord.objects.filter(
            faculty=faculty
        )
        
        if start_date and end_date:
            teaching_records = teaching_records.filter(
                start_time__date__gte=start_date,
                end_time__date__lte=end_date
            )
        
        # Variable to store total salary
        total_salary = 0
        
        # Add rows for each teaching record
        for record in teaching_records:
            # Calculate total salary for the faculty
            total_salary += record.salary
            
            # Convert start_time and end_time to naive datetime (remove timezone)
            start_time_naive = record.start_time.replace(tzinfo=None) if record.start_time else None
            end_time_naive = record.end_time.replace(tzinfo=None) if record.end_time else None
            
            # Append the record data to the sheet
            row_data = [
                record.subject.name,
                record.course.name,
                record.topic_taught,
                start_time_naive,
                end_time_naive,
                record.description,
                record.salary
            ]
            sheet.append(row_data)
        
        # Add total salary at the end of the sheet
        sheet.append(['', '', '', '', '', 'Total Salary', total_salary])
    
    # Remove the default sheet created by openpyxl
    workbook.remove(workbook.active)
    
    # Create an HTTP response with the Excel file as an attachment
    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )
    response['Content-Disposition'] = 'attachment; filename="faculty_data.xlsx"'
    
    # Save the workbook to the response object
    workbook.save(response)
    
    # Return the response
    return response


def error_404(request, exception):
    # Render a custom 404 error page
    return render(request, 'error.html', {}, status=404)

def error_500(request):
    # Render a custom 500 error page (optional)
    return render(request, 'error.html', {}, status=500)